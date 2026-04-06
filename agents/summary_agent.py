import asyncio
import time
import uuid
import re
from typing import Dict, Optional
import logging
import os
import markdown
from docx import Document
from docx.shared import Inches

from agents.fundamental_agent import fundamental_agent
from agents.technical_agent import technical_agent
from agents.market_analysis_agent import market_analysis_agent
from backend.model import model_manager

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class SummaryDecisionAgent:
    def __init__(self):
        self.reports_dir = "reports"
        os.makedirs(self.reports_dir, exist_ok=True)
    
    async def _get_analysis_data(self, symbol: str) -> tuple:
        """获取各类分析数据（基本面、技术面、市场分析）"""
        # 并发获取各类分析数据，提高效率
        fundamental_task = fundamental_agent.analyze_company(symbol)
        technical_task = technical_agent.analyze_stock(symbol)
        market_task = market_analysis_agent.analyze_market()
        
        # 等待所有任务完成
        fundamental_result, technical_result, market_result = await asyncio.gather(
            fundamental_task, technical_task, market_task
        )
        
        return fundamental_result, technical_result, market_result
    
    def _safe_get(self, data: Dict, key: str, default: str = '暂无法获取') -> str:
        """安全获取字典值，处理None和空字符串"""
        value = data.get(key, default)
        return value if value and str(value).strip() else default
    
    async def generate_investment_report(self, symbol: str, user_prompt: Optional[str] = None) -> Dict:
        """生成投资报告"""
        try:
            # 1. 获取各类分析数据
            fundamental_result, technical_result, market_result = await self._get_analysis_data(symbol)
            
            # 2. 构建汇总分析提示词
            prompt = self._build_summary_prompt(
                symbol, 
                fundamental_result, 
                technical_result, 
                market_result, 
                user_prompt
            )
            
            # 5. 调用模型生成投资报告
            report_content = await model_manager.async_generate_response(prompt)
            
            # 6. 生成报告ID
            report_id = str(uuid.uuid4())
            
            # 7. 保存Markdown格式报告
            md_path = os.path.join(self.reports_dir, f"{report_id}.md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(report_content)
            
            # 8. 生成Word格式报告
            docx_path = os.path.join(self.reports_dir, f"{report_id}.docx")
            self._convert_md_to_docx(report_content, docx_path)
            
            # 9. 解析报告内容，提取关键信息
            report_info = self._parse_report(report_content)
            report_info["report_id"] = report_id
            report_info["symbol"] = symbol
            report_info["generated_at"] = int(time.time())
            
            return report_info
            
        except Exception as e:
            logger.error(f"生成投资报告失败: {e}", exc_info=True)
            return {
                "report_id": "",
                "symbol": "",
                "rating": "",
                "target_price": "",
                "risk_level": "",
                "summary": "",
                "generated_at": 0
            }
    
    def _build_summary_prompt(
        self, 
        symbol: str, 
        fundamental: Dict, 
        technical: Dict, 
        market: Dict, 
        user_prompt: Optional[str]
    ) -> str:
        """构建汇总分析提示词"""
        prompt = f"""作为一名专业的投资分析师，请根据以下可获取的公司基本面分析、技术面分析和市场分析，生成一份完整的投资报告：

股票代码：{symbol}

【基本面分析】
估值分析：{fundamental['valuation'] if fundamental['valuation'] else '暂无法获取'}
成长性分析：{fundamental['growth'] if fundamental['growth'] else '暂无法获取'}
财务健康度：{fundamental['financial_health'] if fundamental['financial_health'] else '暂无法获取'}
综合评价：{fundamental['summary'] if fundamental['summary'] else '暂无法获取'}

【估值指标】
PE-TTM：{fundamental.get('pe_ttm', '暂无法获取')}
PB：{fundamental.get('pb', '暂无法获取')}
估值状态：{fundamental.get('valuation_status', '暂无法获取')}
估值说明：{fundamental.get('valuation_detail', '暂无法获取')}

【技术面分析】
技术面评分：{technical['score']:.1f}/10分
趋势分析：{technical['trend_analysis'] if technical['trend_analysis'] else '暂无法获取'}
动量分析：{technical['momentum_analysis'] if technical['momentum_analysis'] else '暂无法获取'}
支撑阻力：{technical['support_resistance'] if 'support_resistance' in technical and technical['support_resistance'] else '暂无法获取'}
短期预测：{technical['short_term_prediction'] if technical['short_term_prediction'] else '暂无法获取'}

【市场分析】
大盘综述：{market.get('market_summary', '暂无法获取')}
强势板块：{', '.join(market.get('strong_sectors', [])) if market.get('strong_sectors', []) else '暂无法获取'}
板块轮动：{market.get('rotation_analysis', '暂无法获取')}
明日板块预测：{market.get('tomorrow_prediction', '暂无法获取')}

"""
        
        if user_prompt:
            prompt += f"\n【用户特殊要求】\n{user_prompt}\n"
        
        prompt += """\n请按照以下格式输出投资报告：

# 投资报告：[股票代码]

## 1. 公司概况
- 公司名称：
- 所属行业：
- 主营业务：

## 2. 投资评级
- 综合评级：（强烈推荐/推荐/中性/谨慎/卖出）
- 目标价格：
- 风险等级：（低/中/高）

## 3. 投资要点
- 要点1：
- 要点2：
- 要点3：

## 4. 详细分析
### 4.1 基本面分析
[详细内容]

### 4.2 技术面分析
[详细内容]

### 4.3 市场环境分析
[详细内容]

## 5. 风险提示
- 风险1：
- 风险2：
- 风险3：

## 6. 投资建议
[详细建议]

请使用专业、客观的语言，确保报告结构清晰、内容完整、分析深入。

【内容要求】
1. **估值分析**：基于PE-TTM和PB数据，结合行业平均水平，分析公司的估值合理性
2. **财务健康度**：分析公司的资产负债率、流动比率等财务指标，评估财务风险
3. **成长性分析**：基于营收增长率、净利润增长率等指标，分析公司的成长潜力
4. **技术面分析**：结合技术指标数据，分析股票的趋势、动量和支撑阻力
5. **市场环境分析**：分析大盘走势和板块轮动情况，评估市场环境对个股的影响
6. **风险提示**：明确指出投资风险，包括行业风险、公司特定风险等
7. **投资建议**：基于综合分析，提供具体的投资策略和操作建议

【数据使用要求】
1. 优先使用提供的具体数据，如PE-TTM、PB、技术面评分等
2. 对于缺失的数据，明确说明"数据缺失"，避免无根据的猜测
3. 结合多个数据点进行综合分析，避免单一指标的片面判断
4. 引用具体数据时，确保数据来源明确，如"根据基本面分析数据"等
"""
        
        return prompt
    
    def _parse_report(self, report_content: str) -> Dict:
        """解析报告内容，提取关键信息"""
        result = {
            "rating": "",
            "target_price": "",
            "risk_level": "",
            "summary": ""
        }
        
        # 提取评级
        if "综合评级：" in report_content:
            rating_start = report_content.find("综合评级：") + 6
            rating_end = report_content.find("\n", rating_start)
            if rating_end != -1:
                result["rating"] = report_content[rating_start:rating_end].strip()
        
        # 提取目标价格
        if "目标价格：" in report_content:
            tp_start = report_content.find("目标价格：") + 6
            tp_end = report_content.find("\n", tp_start)
            if tp_end != -1:
                result["target_price"] = report_content[tp_start:tp_end].strip()
        
        # 提取风险等级
        if "风险等级：" in report_content:
            risk_start = report_content.find("风险等级：") + 6
            risk_end = report_content.find("\n", risk_start)
            if risk_end != -1:
                result["risk_level"] = report_content[risk_start:risk_end].strip()
        
        # 提取摘要（投资要点）
        if "## 3. 投资要点" in report_content:
            summary_start = report_content.find("## 3. 投资要点") + len("## 3. 投资要点")
            summary_end = report_content.find("## 4.", summary_start)
            if summary_end != -1:
                result["summary"] = report_content[summary_start:summary_end].strip()
        
        return result
    
    def _convert_md_to_docx(self, md_content: str, output_path: str) -> None:
        """将Markdown转换为Word文档"""
        try:
            # 创建文档
            doc = Document()
            
            # 简单转换Markdown内容
            lines = md_content.split("\n")
            current_heading_level = 0
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # 处理标题
                if line.startswith("#"):
                    heading_level = len(line.split(" ")[0])
                    title_text = line[heading_level:].strip()
                    
                    if heading_level == 1:
                        doc.add_heading(title_text, level=1)
                    elif heading_level == 2:
                        doc.add_heading(title_text, level=2)
                    elif heading_level == 3:
                        doc.add_heading(title_text, level=3)
                    elif heading_level == 4:
                        doc.add_heading(title_text, level=4)
                    
                    current_heading_level = heading_level
                
                # 处理列表项
                elif line.startswith("-"):
                    doc.add_paragraph(line[1:].strip(), style='ListBullet')
                
                # 处理正文
                else:
                    doc.add_paragraph(line)
            
            # 保存文档
            doc.save(output_path)
            logger.info(f"Word文档已保存: {output_path}")
            
        except Exception as e:
            logger.error(f"转换Markdown到Word失败: {e}")
    
    def get_report_path(self, report_id: str, format: str = "docx") -> Optional[str]:
        """获取报告文件路径"""
        if format == "docx":
            file_path = os.path.join(self.reports_dir, f"{report_id}.docx")
        elif format == "md":
            file_path = os.path.join(self.reports_dir, f"{report_id}.md")
        else:
            return None
        
        return file_path if os.path.exists(file_path) else None
    
    async def generate_quick_analysis(self, symbol: str, user_prompt: Optional[str] = None) -> str:
        """生成快速投资分析，用于直接在对话栏回答用户"""
        try:
            # 1. 获取各类分析数据
            fundamental_result, technical_result, market_result = await self._get_analysis_data(symbol)
            
            # 2. 构建快速分析提示词
            # 获取股票名称
            stock_name = fundamental_result.get('stock_name', '') or symbol
            
            prompt = f"# 投资分析任务\n\n"
            prompt += f"**分析对象**：{symbol}({stock_name})\n\n"
            
            # 整理并结构化可用数据
            available_data = {}
            
            # 基本面数据
            if any(fundamental_result.values()):
                available_data['基本面'] = {
                    '估值分析': fundamental_result.get('valuation'),
                    '成长性分析': fundamental_result.get('growth'),
                    '财务健康度': fundamental_result.get('financial_health'),
                    '综合评价': fundamental_result.get('summary'),
                    'PE-TTM': fundamental_result.get('pe_ttm'),
                    'PB': fundamental_result.get('pb'),
                    '估值状态': fundamental_result.get('valuation_status'),
                    '估值说明': fundamental_result.get('valuation_detail')
                }
            
            # 技术面数据
            if any(technical_result.values()):
                available_data['技术面'] = {
                    '评分': technical_result.get('score'),
                    '趋势分析': technical_result.get('trend_analysis'),
                    '动量分析': technical_result.get('momentum_analysis'),
                    '短期预测': technical_result.get('short_term_prediction')
                }
            
            # 市场分析数据
            if market_result.get('market_summary') or market_result.get('strong_sectors') or market_result.get('rotation_analysis') or market_result.get('tomorrow_prediction') or market_result.get('overseas_indices'):
                available_data['市场分析'] = {
                    '大盘综述': market_result.get('market_summary'),
                    '强势板块': market_result.get('strong_sectors'),
                    '板块轮动': market_result.get('rotation_analysis'),
                    '明日板块预测': market_result.get('tomorrow_prediction'),
                    '海外指数': market_result.get('overseas_indices')
                }
            
            # 构建数据可用性说明
            prompt += "## 可用数据\n"
            if available_data:
                # 按优先级列出可用数据类型
                for data_type, data_content in available_data.items():
                    valid_data = {k: v for k, v in data_content.items() if v}
                    if valid_data:
                        prompt += f"- **{data_type}数据**：{', '.join(valid_data.keys())}\n"
            else:
                prompt += "- 暂无可用的分析数据\n"
            prompt += "\n"
            
            # 详细列出所有可用的具体数据
            if available_data.get('基本面'):
                prompt += "### 基本面详细数据\n"
                for key, value in available_data['基本面'].items():
                    if value:
                        if key in ['PE-TTM', 'PB', '估值状态', '估值说明']:
                            prompt += f"- **{key}**：{value}\n"
                        else:
                            prompt += f"- **{key}**：{value}\n"
                prompt += "\n"
            
            if available_data.get('技术面'):
                prompt += "### 技术面详细数据\n"
                for key, value in available_data['技术面'].items():
                    if value:
                        if key == '评分' and isinstance(value, (int, float)):
                            prompt += f"- **{key}**：{value:.1f}/10分\n"
                        elif key == '强势板块' and isinstance(value, list):
                            prompt += f"- **{key}**：{', '.join(value)}\n"
                        else:
                            prompt += f"- **{key}**：{value}\n"
                prompt += "\n"
            
            if available_data.get('市场分析'):
                prompt += "### 市场分析详细数据\n"
                for key, value in available_data['市场分析'].items():
                    if value:
                        if isinstance(value, list):
                            prompt += f"- **{key}**：{', '.join(value)}\n"
                        elif isinstance(value, dict) and key == '海外指数':
                            # 特殊处理海外指数数据
                            prompt += f"- **{key}**：\n"
                            for index_symbol, index_data in value.items():
                                if isinstance(index_data, dict):
                                    name = index_data.get('name', index_symbol)
                                    price = index_data.get('price', 'N/A')
                                    change = index_data.get('change_percent', 'N/A')
                                    prompt += f"  - {name}：{price} ({change}%)\n"
                        else:
                            prompt += f"- **{key}**：{value}\n"
                prompt += "\n"
            
            if user_prompt:
                prompt += f"## 用户特殊要求\n{user_prompt}\n\n"
            
            # 任务指令
            prompt += "## 任务要求\n"
            prompt += "1. **分析范围**：基于上述提供的数据，不得引入外部信息或主观臆断\n"
            prompt += "2. **内容要求**：完全围绕股票投资分析，不得包含任何无关内容\n"
            prompt += "3. **数据使用**：优先使用数据量最丰富的数据类型，对于缺失数据明确说明\n"
            prompt += "4. **权重分配**：基本面 > 技术面 > 市场分析\n"
            prompt += "5. **分析深度**：结合具体数据进行深入分析，避免泛泛而谈\n"
            prompt += "6. **投资建议**：提供具体的投资策略和操作建议\n"
            prompt += "7. **估值分析**：基于PE-TTM和PB数据，分析公司的估值合理性\n"
            prompt += "8. **财务健康**：分析公司的财务指标，评估财务风险\n"
            prompt += "9. **成长性**：基于营收和净利润增长率，分析公司的成长潜力\n"
            prompt += "10. **技术面**：结合技术指标数据，分析股票的趋势和动量\n\n"
            
            # 格式要求
            prompt += "## 格式要求\n"
            prompt += "- 语言：专业、客观、清晰易懂\n"
            prompt += "- 长度：300-400字\n"
            prompt += "- 结构：层次分明，重点突出\n"
            prompt += "- 避免：复杂术语、无关内容、无根据预测\n"
            prompt += "- 数据引用：明确引用提供的数据，如PE-TTM、PB等\n"
            prompt += "- 分析逻辑：从数据到结论，逻辑清晰，论证充分\n"
            prompt += "- 投资建议：具体、可操作，包含仓位建议和操作时机\n\n"
            
            # 生成数据可用性说明
            data_availability = []
            if available_data.get('基本面'):
                data_availability.append("基本面数据")
            if available_data.get('技术面'):
                data_availability.append("技术面数据")
            if available_data.get('市场分析'):
                data_availability.append("市场分析数据")
            
            if data_availability:
                data_availability_str = "、".join(data_availability)
            else:
                data_availability_str = "暂无数据"
            
            # 开始分析指令
            prompt += f"请直接开始分析，严格按照以下格式输出，不要添加任何开场白或引言：\n\n"
            prompt += f"{symbol} {stock_name} 投资分析报告\n\n"
            prompt += f"1. 股票概况\n"
            prompt += f"- 代码：{symbol}\n"
            prompt += f"- 名称：{stock_name}\n"
            prompt += f"- 数据可用性：{data_availability_str}（注：基本面数据中可能缺失当前精确股价及部分历史对比数据）\n\n"
            prompt += f"2. 核心分析\n"
            prompt += f"基本面分析\n"
            prompt += f"- 财务健康：基于财务健康度数据进行分析，说明公司财务结构和偿债能力。\n"
            prompt += f"- 盈利能力：分析公司的盈利情况，包括净利润、ROE等指标。\n"
            prompt += f"- 估值水平：PE-TTM {fundamental_result.get('pe_ttm', '暂无数据')}，PB {fundamental_result.get('pb', '暂无数据')}，{fundamental_result.get('valuation_detail', '暂无估值说明')}。\n\n"
            prompt += f"技术面分析\n"
            prompt += f"- 趋势：分析股票的趋势，包括均线排列等。\n"
            prompt += f"- 动量：基于MACD、KDJ等指标，分析股票的动量状态。\n"
            prompt += f"- 评分：{technical_result.get('score', 'N/A')}/10\n\n"
            prompt += f"市场环境分析\n"
            prompt += f"- 大盘情绪：分析大盘的整体表现和市场情绪。\n"
            prompt += f"- 板块轮动：分析板块的轮动情况和资金流向。\n\n"
            prompt += f"3. 综合评估\n"
            prompt += f"优势\n"
            prompt += f"- 列出个股的主要优势和潜在利好因素。\n"
            prompt += f"风险\n"
            prompt += f"- 明确指出投资风险和不确定性因素。\n"
            prompt += f"总体判断\n"
            prompt += f"基于以上分析，给出对个股的总体评价。\n\n"
            prompt += f"4. 投资建议\n"
            prompt += f"- 短期策略：提供具体的短期操作建议，如观望、试探性建仓、逢高减仓等，并说明理由。\n"
            prompt += f"- 中长期策略：基于基本面和行业趋势，给出中长期投资建议。\n"
            prompt += f"- 仓位建议：根据风险等级，建议合理的仓位配置。\n\n"
            prompt += f"5. 操作提示\n"
            prompt += f"- 如需更详细报告，请输入\"生成投资报告\"。\n"
            prompt += f"- 如需下载Word文档，请输入\"下载报告\"。\n"
            prompt += f"- 以上分析基于当前数据，不构成投资建议，决策需结合实时市场情况。\n"

            
            # 5. 调用模型生成快速分析
            logger.debug(f"传递给模型的完整提示词: {prompt}")
            analysis_content = await model_manager.async_generate_response(prompt)
            logger.debug(f"模型返回的原始内容: {analysis_content}")
            
            # 检查模型返回是否包含错误信息
            error_patterns = ["模型未加载成功", "PyTorch不可用", "无法生成响应"]
            for pattern in error_patterns:
                if pattern in analysis_content:
                    logger.warning(f"模型生成失败: {analysis_content}")
                    return f"分析服务暂时不可用，请稍后重试。\n\n错误信息: {analysis_content}"
            
            # 6. 验证生成内容的相关性
            def validate_response_relevance(response, symbol, stock_name, available_data):
                """验证生成内容的相关性"""
                # 处理可能的股票名称为空的情况
                has_stock_info = False
                
                # 调试日志
                logger.info(f"验证相关性: 响应='{response[:100]}...', 股票代码='{symbol}', 股票名称='{stock_name}'")
                
                # 检查是否包含错误的股票代码（如000470而不是002460）
                wrong_code_patterns = [r'000470', r'中联重工']
                for pattern in wrong_code_patterns:
                    if re.search(pattern, response):
                        logger.warning(f"检测到错误的股票代码或名称: {pattern}")
                        return False, f"内容包含错误的股票代码或名称：{pattern}"
                
                # 使用更精确的正则表达式检查股票代码，支持各种格式
                if symbol:
                    # 匹配股票代码，支持：002460、(002460)、002460)、(002460 等格式
                    symbol_pattern = re.compile(r'[\s\(\[]*{}[\s\)\]]*'.format(re.escape(symbol)))
                    if symbol_pattern.search(response):
                        has_stock_info = True
                        logger.info(f"匹配到股票代码: {symbol}")
                
                # 如果有股票名称，检查是否包含股票名称
                if not has_stock_info and stock_name:
                    stock_name_stripped = stock_name.strip()
                    if stock_name_stripped:
                        # 匹配股票名称，忽略大小写和部分匹配
                        name_pattern = re.compile(re.escape(stock_name_stripped), re.IGNORECASE)
                        if name_pattern.search(response):
                            has_stock_info = True
                            logger.debug(f"匹配到股票名称: {stock_name_stripped}")
                    
                    # 检查是否包含股票名称的简称
                    if not has_stock_info:
                        # 提取股票名称的简称（如"江西赣锋锂业集团股份有限公司"的简称为"赣锋锂业"）
                        stock_name_parts = stock_name_stripped.split()
                        for part in stock_name_parts:
                            if len(part) > 2:  # 只检查长度大于2的部分
                                if part in response:
                                    has_stock_info = True
                                    logger.debug(f"匹配到股票名称简称: {part}")
                                    break

                # 对于包含股票信息的内容，更加宽松
                if has_stock_info:
                    # 检查是否包含不相关的典型内容
                    irrelevant_patterns = ["账号", "权利", "服务", "祝福", "财务知识", "个人所得税", "登录", "注册",
                                     "保持持续努力", "无需害怕犯错误", "坚守真诚", "初心", "大胆去追逐"]
                    for pattern in irrelevant_patterns:
                        if pattern in response:
                            logger.warning(f"内容包含不相关信息: {pattern}")
                            return False, f"内容包含不相关信息：{pattern}"
                    
                    # 检查是否使用了技术指标数据（如果提供了技术面数据）
                    if available_data.get('技术面'):
                        technical_keywords = ["MA", "MACD", "BOLL", "KDJ", "均线", "金叉", "死叉", "趋势", "支撑", "阻力", "动量"]
                        has_technical_analysis = any(keyword in response for keyword in technical_keywords)
                        if not has_technical_analysis:
                            logger.warning("内容未使用技术指标数据")
                            return False, "内容未使用技术指标数据"
                    
                    # 如果包含股票信息，即使没有足够的关键词也认为相关
                    logger.info("内容包含股票信息，认为相关")
                    return True, "内容相关"

                # 如果没有包含股票信息，检查是否有足够的投资分析关键词
                relevant_keywords = ["基本面", "技术面", "投资建议", "优势", "风险", "趋势", "分析", "报告", "数据",
                                   "估值", "增长", "财务", "健康", "收益", "风险", "前景", "评分", "市场", "板块"]
                keyword_count = sum(1 for keyword in relevant_keywords if keyword in response)
                
                logger.debug(f"未匹配到股票信息，关键词数量: {keyword_count}")
                
                if keyword_count < 1:
                    # 检查是否包含提示词中的分析结构要求
                    structure_keywords = ["股票概况", "核心分析", "综合评估", "投资建议", "操作提示"]
                    structure_count = sum(1 for keyword in structure_keywords if keyword in response)
                    
                    if structure_count < 2:
                        return False, "内容未包含股票代码或名称且缺乏投资分析核心要素"
                    else:
                        # 包含分析结构，认为相关
                        logger.debug("内容包含分析结构，认为相关")
                        return True, "内容包含分析结构"
                
                return True, "内容相关"
            
            # 进行相关性验证
            is_relevant, reason = validate_response_relevance(analysis_content, symbol, stock_name, available_data)
            
            if not is_relevant:
                logger.warning(f"生成内容不相关: {reason}")
                # 如果内容不相关，生成引导性响应
                return f"未能为{symbol}({stock_name})生成有效的投资分析。\n\n请确保您提供了正确的股票代码或名称，系统将基于准确的股票信息为您生成专业分析。\n\n您可以尝试：\n1. 提供完整的股票代码（如：002460）\n2. 提供准确的股票名称（如：赣锋锂业）\n3. 稍后重试，系统将重新获取最新数据"
            
            return analysis_content
            
        except Exception as e:
            logger.error(f"生成快速投资分析失败: {e}", exc_info=True)
            return f"获取股票 {symbol} 的分析信息失败，请稍后重试。错误原因: {str(e)[:100]}..."
    
    async def generate_quick_advice(self, symbol: str) -> str:
        """生成快速投资建议"""
        try:
            # 并发获取基本面和技术面分析
            fundamental_task = fundamental_agent.analyze_company(symbol)
            technical_task = technical_agent.analyze_stock(symbol)
            
            fundamental, technical = await asyncio.gather(fundamental_task, technical_task)
            
            # 构建快速建议提示词，考虑数据可能缺失的情况
            prompt = f"""请根据以下公司的基本面和技术面分析（如数据缺失请明确说明），提供一句简洁的投资建议：

股票代码：{symbol}

基本面分析：{'数据缺失' if not fundamental.get('summary') else fundamental['summary']}
技术面分析：技术面评分 {'数据缺失' if not technical.get('score') else f"{technical['score']:.1f}/10分"}，{'数据缺失' if not technical.get('short_term_prediction') else technical['short_term_prediction']}

分析要求：
1. 基于提供的所有可用数据进行分析
2. 如果某些数据缺失，请在理由中明确说明
3. 不要使用'数据限制'等模糊词汇，要具体说明哪些数据缺失
4. 保持建议简洁明了，符合投资逻辑

建议格式：[股票代码] - [买入/持有/卖出] - [简短理由]
"""
            
            # 调用模型生成快速建议
            advice = await model_manager.async_generate_response(prompt)
            return advice
            
        except Exception as e:
            logger.error(f"生成快速建议失败: {e}")
            return f"{symbol} - 中性 - 分析数据不可用"
    
    async def calculate_combined_score(self, symbol: str) -> Dict:
        """计算综合评分，整合三个分析面的结果
        
        Returns:
            Dict: 包含各分析面评分和综合评分的字典
        """
        try:
            # 并发获取各类分析数据
            fundamental_task = fundamental_agent.analyze_company(symbol)
            technical_task = technical_agent.analyze_stock(symbol)
            market_task = market_analysis_agent.analyze_market()
            
            fundamental_result, technical_result, market_result = await asyncio.gather(
                fundamental_task, technical_task, market_task
            )
            
            # 计算基本面评分
            fundamental_score = self._calculate_fundamental_score(fundamental_result)
            
            # 技术面评分直接使用technical_result中的score
            technical_score = technical_result.get('score', 5.0)
            
            # 计算市场分析评分
            market_score = self._calculate_market_score(market_result, symbol)
            
            # 应用权重
            weights = {
                'fundamental': 0.5,    # 基本面：50%
                'technical': 0.3,       # 技术面：30%
                'market': 0.2           # 市场分析：20%
            }
            
            combined_score = (
                fundamental_score * weights['fundamental'] +
                technical_score * weights['technical'] +
                market_score * weights['market']
            )
            
            # 确定投资结论
            if combined_score >= 7:
                conclusion = "强烈关注"
            elif combined_score >= 5:
                conclusion = "中性"
            else:
                conclusion = "谨慎"
            
            # 例外调整
            exceptions = []
            if technical_score < 3 and fundamental_score > 7:
                exceptions.append("基本面好但短期技术走坏，可等待买点")
            
            # 计算市场情绪
            market_sentiment = self._calculate_market_sentiment(market_result)
            if market_sentiment == "冰点" and fundamental_score >= 6:
                exceptions.append("市场过度悲观，可能存在错杀机会")
            
            return {
                'fundamental_score': fundamental_score,
                'technical_score': technical_score,
                'market_score': market_score,
                'combined_score': combined_score,
                'conclusion': conclusion,
                'exceptions': exceptions
            }
            
        except Exception as e:
            logger.error(f"计算综合评分失败: {e}")
            return {
                'fundamental_score': 5.0,
                'technical_score': 5.0,
                'market_score': 5.0,
                'combined_score': 5.0,
                'conclusion': "中性",
                'exceptions': []
            }
    
    def _calculate_fundamental_score(self, fundamental_result: Dict) -> float:
        """计算基本面评分（0-10分）"""
        try:
            # 估值分析评分（1-3分）
            valuation = fundamental_result.get('valuation', '')
            valuation_score = 2.0  # 默认中性
            if "低估" in valuation:
                valuation_score = 3.0
            elif "高估" in valuation:
                valuation_score = 1.0
            
            # 成长性分析评分（1-3分）
            growth = fundamental_result.get('growth', '')
            growth_score = 2.0  # 默认中性
            if "高增长" in growth or "增长" in growth:
                growth_score = 3.0
            elif "下降" in growth or "负增长" in growth:
                growth_score = 1.0
            
            # 财务健康度评分（1-3分）
            financial_health = fundamental_result.get('financial_health', '')
            financial_score = 2.0  # 默认中性
            if "健康" in financial_health:
                financial_score = 3.0
            elif "风险" in financial_health or "预警" in financial_health:
                financial_score = 1.0
            
            # 综合评分
            total_score = (valuation_score + growth_score + financial_score) / 3 * 10 / 3
            return min(10.0, max(0.0, total_score))
            
        except Exception as e:
            logger.error(f"计算基本面评分失败: {e}")
            return 5.0
    
    def _calculate_market_score(self, market_result: Dict, symbol: str) -> float:
        """计算市场分析评分（0-10分）"""
        try:
            score = 5.0  # 默认中性
            
            # 分析大盘综述
            market_summary = market_result.get('market_summary', '')
            if "上涨" in market_summary:
                score += 1.0
            elif "下跌" in market_summary:
                score -= 1.0
            
            # 分析强势板块
            strong_sectors = market_result.get('strong_sectors', [])
            if strong_sectors:
                # 假设股票所属行业在强势板块中
                # 这里简化处理，实际应该根据股票所属行业判断
                score += 1.0
            
            # 分析板块轮动
            rotation_analysis = market_result.get('rotation_analysis', '')
            if "轮动明显" in rotation_analysis or "热点" in rotation_analysis:
                score += 0.5
            
            # 分析明日板块预测
            tomorrow_prediction = market_result.get('tomorrow_prediction', '')
            if "领涨" in tomorrow_prediction:
                score += 0.5
            
            return min(10.0, max(0.0, score))
            
        except Exception as e:
            logger.error(f"计算市场评分失败: {e}")
            return 5.0
    
    def _calculate_market_sentiment(self, market_result: Dict) -> str:
        """计算市场情绪"""
        try:
            market_summary = market_result.get('market_summary', '')
            
            # 简单判断市场情绪
            if "上涨" in market_summary and "活跃" in market_summary:
                return "火热"
            elif "下跌" in market_summary and "谨慎" in market_summary:
                return "冰点"
            else:
                return "正常"
            
        except Exception as e:
            logger.error(f"计算市场情绪失败: {e}")
            return "正常"

# 全局汇总决策Agent实例
summary_agent = SummaryDecisionAgent()
